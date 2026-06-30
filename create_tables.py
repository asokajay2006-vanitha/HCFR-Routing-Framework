from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# --------------------------------------------------------
# Function
# --------------------------------------------------------

def add_table(title, headers, data):

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr = table.rows[0].cells

    for i, h in enumerate(headers):
        hdr[i].text = h

    for row in data:

        cells = table.add_row().cells

        for j, item in enumerate(row):
            cells[j].text = str(item)

    doc.add_paragraph()


# --------------------------------------------------------
# Table 2
# --------------------------------------------------------

headers = [
    "Method",
    "Accuracy (%)",
    "Precision (%)",
    "Recall (%)",
    "F1-Score"
]

data = [

    ["ML-IDS",91.4,89.7,88.9,0.89],
    ["AAR",93.2,91.8,90.6,0.91],
    ["HCFR\n(Proposed)",97.1,96.4,95.8,0.96]

]

add_table(
    "Table 2. Detection Performance Comparison",
    headers,
    data
)

# --------------------------------------------------------
# Table 3
# --------------------------------------------------------

headers = [
    "Configuration",
    "Detection Accuracy (%)",
    "PDR (%)",
    "Delay (ms)",
    "Throughput Stability"
]

data = [

["Full HCFR",97.1,96.8,97,0.91],

["HCFR without Sparse Feature Selection",
94.3,93.7,108,0.86],

["HCFR without Hierarchical SNN",
92.8,92.1,116,0.83],

["HCFR without DAG Optimization",
91.9,90.8,121,0.81],

["HCFR without Fractional Routing",
90.6,89.7,128,0.79]

]

add_table(
"Table 3. Ablation Analysis of HCFR Components",
headers,
data
)

# --------------------------------------------------------
# Table 4
# --------------------------------------------------------

headers = [
"Method",
"PDR (%)",
"Packet Loss (%)"
]

data = [

["SPR",84.6,15.4],
["LAR",88.9,11.1],
["AAR",91.7,8.3],
["HCFR",96.8,3.2]

]

add_table(
"Table 4. Packet Delivery Ratio and Packet Loss",
headers,
data
)

# --------------------------------------------------------
# Table 5
# --------------------------------------------------------

headers = [
"Method",
"Delay (ms)"
]

data = [

["SPR",142],
["LAR",128],
["AAR",119],
["HCFR",97]

]

add_table(
"Table 5. Average End-to-End Delay",
headers,
data
)

# --------------------------------------------------------
# Table 6
# --------------------------------------------------------

headers = [
"Method",
"Stability Index"
]

data = [

["SPR",0.71],
["ML-IDS",0.78],
["AAR",0.83],
["HCFR",0.91]

]

add_table(
"Table 6. Throughput Stability Index",
headers,
data
)

doc.save("HCFR_Tables.docx")

print("HCFR_Tables.docx Created Successfully")